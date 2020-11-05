import sys
import re
import docx
from snapgene_reader import snapgene_file_to_dict
from restrictiondiction import rsites
import docx2txt

def errchk(line):
        search= "[^ACGTacgt\s+]"
        bad=re.search(search,line)
        if bad:
            return "F",bad
        else:
            return "True"

def optchk(opt,input):
    position=0
    for item in input:
        if item==opt:
            if opt=='-r':
                return input[position+1:position+3]
            else:
                return input[position+1]
        position+=1

def comchk(line):
    search='-\w -'
    wrongorder=re.search(search,line)
    if wrongorder:
        return 'FAIL'
    else:
        return 'PASS'

def parachk(doc):
    empty=0
    for para in doc:
        if para.text== "":
            empty+=1
    if empty != 0:
        return "FAIL"
    else: return "PASS"

#importing all the data
NO=0
errcount=0
counter=3
bad=0
command=sys.argv[:]
commandstring=" ".join(command)
commandorderchk=comchk(commandstring)
bcount=command.count("-b")
ocount=command.count("-o")
icount=command.count("-i")
rcount=command.count("-r")
commandcount=len(command)
insert=docx.Document(optchk('-i',command))
insert=insert.paragraphs
format=parachk(insert)
if format== "FAIL":
    while insert[-1].text=="":
        insert=insert[:-1]
    check2=parachk(insert)
    if check2 == "FAIL":
        errcount+=1
        NO+=1
        bad+=1
        print("ERROR: FASTA file was not formatted correctly.\nSeparate headers and sequences by only one 'ENTER'")
if bcount==1 and ocount==1 and icount==1 and rcount==1 and commandcount==10 and commandorderchk=='PASS':
    outdoc=docx.Document()
    bb=optchk('-b',command)
    bb=snapgene_file_to_dict(bb)
    bb=bb['seq']
    bb=bb.upper()
    totbackbone=bb
    restrictions=optchk('-r',command)
    restriction1=restrictions[0]
    restriction1=restriction1.upper()
    restriction2=restrictions[1]
    restriction2=restriction2.upper()
    z=optchk('-o',command)
    if z[-5:]!='.docx':
        z=z+'.docx'
    with open (z,"w") as outfile:
        if restriction1 not in rsites or restriction2 not in rsites:
            print("ERROR: Invalid Restriction Site(s).")
        else:
            echeck=errchk(totbackbone)
            if echeck[0]== "F":
                print("ERROR: Invalid Character(s) in Backbone Sequence.",echeck[1:])
                NO+=1
            else:
                for line in insert:
                    if counter%2==1:
                        outtitle=line.text.replace('pMAX','mRNA')
                        outtitle=outtitle+"-dTomato Ligation Sequence"
                        outdoc.add_paragraph(outtitle)
                    else:
                        errorcheck=errchk(line.text)
                        if errorcheck[0]=="F" and bad==0:
                            print("ERROR: Invalid Character(s) in Insert Sequence.",errorcheck[1:])
                            NO+=1
                            errcount+=1
                        else:

                            insert="CGCGT"+line.text+"A"
                            insert=insert.lower()
                            insert=insert.strip()
                            insert=insert.replace(" ","")
                            insert=insert.replace("\t","")
                            site=rsites[restriction1]
                            fsite=site.replace('/',"")
                            firstcut=re.sub(fsite,site,totbackbone)
                            site2=rsites[restriction2]
                            fsite2=site2.replace('/',"")
                            secondcut=re.sub(fsite2,site2,firstcut)
                            countsecondcut=secondcut.count("/")
                            if countsecondcut >2 and errcount ==0:
                                ask=input("WARNING: Your Restriction enzymes make more than two cuts in your backbone. Continue? y/n: ")
                                if ask=="n":
                                    NO+=1
                                    print('\nProgram Terminated by User.')
                                    break
                                if ask =="y":
                                    errcount+=1
                            doubledigest= "\/.+\/"
                            ligation=re.sub(doubledigest,insert,secondcut)
                            outdoc.add_paragraph(ligation)
                    counter+=1
    if NO ==0:
        outpath='OutFiles/'+z
        outdoc.save(outpath)
        #write file as .dna
        txtfile=docx2txt.process(outpath)
        outpath=outpath.replace('.docx','.dna')
        with open(outpath,'w') as outtxtfile:
            outtxtfile.write(txtfile)
        print("\nDone!\nYou can find your new FASTA and SnapGene files in the 'OutFiles' directory located on this drive.")
else:
    print('\nERROR: Improper Command Line Input.\nCorrect usage requires these elements in any order:\n"-b path to backbone.dna"\n"-i path to insert.docx"\n"-o AssignNameForNewOutFiles"\n"-r RestrictionSite#1 RestrictionSite#2"')
    #things left to do
    #1)just some more extensive testing
